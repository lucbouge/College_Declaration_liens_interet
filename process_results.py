import os
import re
from bs4 import BeautifulSoup, Tag, PageElement, NavigableString
from typing import NamedTuple, Optional
import pprint as pp
import pandas as pd
from docx import Document
import datetime
from unidecode import unidecode

## https://github.com/mwilliamson/python-mammoth#readme
import mammoth


original_docx_filename = "2021-11-23_results.docx"
# original_docx_filename = "test.docx"


class Item(NamedTuple):
    firstname: str
    lastname: str
    h1: str
    h2: str
    h3: Optional[str]
    question: str
    answer: str


survey_table = {
    "Données générales": (
        "Votre prénom",
        "Votre nom de famille",
        "Votre adresse électronique",
        "Votre affiliation principale aujourd'hui",
        "Votre adresse professionnelle",
        "Votre adresse personnelle",
        "Votre numéro de téléphone portable",
    ),
    "Activité principale": (
        "Nature de l'activité",
        "Activité",
        "Lieu d’exercice",
        "Début",
        "Fin",
    ),
    "Activité secondaire": (
        "Nature de l'activité",
        "Organisme",
        "Domaine et type de travaux",
        "Rémunération: montant et beneficiaire",
        "Début",
        "Fin",
    ),
    "Activité financée": (
        "Nature de l'activité",
        "Structures et activités des bénéficiaires du financement",
        "Organisme financeur",
        "Pourcentage montant/budget, selon l'échelle ci-dessus",
        "Montant, en ordre de grandeur",
        "Début",
        "Fin",
    ),
    "Participation": {
        "Structure concernée",
        "Type d'investissement",
        "Part dans le capital",
        "Montant détenu",
        "Debut",
        "Fin",
    },
    "Engagement": (
        "Cadre de l'activité",
        "Statut dans ce cadre",
        "Fonctions exercées",
        "Début",
        "Fin",
    ),
    "Autres liens": (
        "Elément ou fait concerné",
        "Montant des sommes perçues si c'est le cas",
        "Commentaires",
        "Début",
        "Fin",
    ),
}

persons = (
    ("Eric", "Arquis"),
    ("Luc", "Bougé"),
    ("Anne", "Guillaume"),
    ("Louise", "Nyssen"),
    ("Marc", "Taillefer"),
    ("Guy", "Wormser"),
    ##
    ("Bruno", "Allard"),
    ("Patrick", "Lemaire"),
    ("Clément", "Léna"),
    ("François", "Massol"),
    ("Rémi", "Mounier"),
    ("Benoît", "Schoefs"),
    ##
    ("Magali", "Boutrais"),
    ("Florence", "Hachez-Leroy"),
    ("Marie-Pierre", "Julien"),
    ("Pierre", "Lurbe"),
    ("Sylvie", "Pittia"),
    ("Dominique", "Valérian"),
    ##
    ("Claire", "Dupas"),
)


################################################################################


def main():
    with open(original_docx_filename, "rb") as cin:
        result = mammoth.convert_to_html(cin)
        text = result.value
        text = "<html><body>" + text + "</body></html>"
    with open("output.html", "w") as cout:
        cout.write(text)

    ##
    soup = BeautifulSoup(text, features="lxml")
    body = soup.body
    assert body is not None
    items = make_items(body=body)
    # pp.pprint(items)
    ##
    df = pd.DataFrame.from_records(items, columns=Item._fields)
    # df.to_excel("results.xlsx", index=False)
    check_df(df)
    ##
    key_to_answer_dict = make_key_to_answer_dict(items)
    existing_persons = set((item.firstname, item.lastname) for item in items)
    new_persons = set(persons) - existing_persons
    key_to_answer_dict = extend_key_to_answer_dict(
        new_persons=new_persons, key_to_answer_dict=key_to_answer_dict
    )
    make_documents(
        key_to_answer_dict=key_to_answer_dict,
        persons=persons,
    )


def extend_key_to_answer_dict(*, new_persons, key_to_answer_dict):
    for (firstname, lastname) in new_persons:
        print(f"New person: {firstname} {lastname}")
        for (k, v) in (("Votre prénom", firstname), ("Votre nom de famille", lastname)):
            key = (
                firstname,
                lastname,
                "Données générales",
                None,
                k,
            )
            assert key not in key_to_answer_dict
            key_to_answer_dict[key] = v
    return key_to_answer_dict


################################################################################


def make_documents(*, key_to_answer_dict, persons):
    document = Document()
    document.add_heading("Déclarations de liens d'intérêt", 0)
    now = datetime.datetime.now()
    now_string = s = now.strftime("%d/%m/%Y, %H:%M")
    document.add_paragraph(f"Date: {now_string}")
    ##
    for (person_nb, (firstname, lastname)) in enumerate(persons, start=1):
        document = Document()
        document.add_heading("Déclarations de liens d'intérêt", 0)
        now = datetime.datetime.now()
        now_string = s = now.strftime("%d/%m/%Y, %H:%M")
        document.add_paragraph(f"Date: {now_string}")
        assert document is not None
        # document.add_page_break()
        print("Personne", (firstname, lastname))
        make_document_person(
            key_to_answer_dict=key_to_answer_dict,
            firstname=firstname,
            lastname=lastname,
            person_nb=person_nb,
            document=document,
        )
        filename = f"declaration_{lastname}_{firstname}.docx"
        filename = unidecode(filename).lower()
        dirname = "Declarations"
        os.makedirs(dirname, exist_ok=True)
        document.core_properties.language = "fr-FR"
        document.save(os.path.join(dirname, filename))


def make_document_person(
    *, key_to_answer_dict, firstname, lastname, person_nb, document
):
    document.add_heading(f"Personne {person_nb}: {lastname}, {firstname}", 1)
    for (block_nb, (block, questions)) in enumerate(survey_table.items(), start=1):
        # print("Bloc", block)
        document.add_heading(f"Bloc {block_nb}: {block}", 2)
        if block == "Données générales":
            make_document_questions(
                firstname=firstname,
                lastname=lastname,
                h2=block,
                h3=None,
                prefix=(person_nb, block_nb, 1),
                questions=questions,
                key_to_answer_dict=key_to_answer_dict,
                document=document,
            )
            continue
        for series_nb in (1, 2, 3, 4):
            series = f"Réponse {series_nb}"
            document.add_heading(f"{series}", 3)
            make_document_questions(
                firstname=firstname,
                lastname=lastname,
                h2=block,
                h3=series,
                prefix=(person_nb, block_nb, series_nb),
                questions=questions,
                key_to_answer_dict=key_to_answer_dict,
                document=document,
            )


def make_document_questions(
    *, firstname, lastname, h2, h3, prefix, questions, key_to_answer_dict, document
):
    for (question_nb, question) in enumerate(questions, start=1):
        answer = get_answer(
            firstname=firstname,
            lastname=lastname,
            h2=h2,
            h3=h3,
            question=question,
            key_to_answer_dict=key_to_answer_dict,
        )
        p = document.add_paragraph(style="List Bullet")
        full_prefix = prefix + (question_nb,)
        prefix_string = ".".join((str(s) for s in full_prefix))
        p.add_run(f"[{prefix_string}] ")
        p.add_run(f"{question}: ").bold = True
        p.add_run(f"{answer}")
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[
        0
    ].text = (
        f"Si le bloc est non vide, j'écris 'OK'\ndans la case de droite pour le valider"
    )
    hdr_cells[1].text = ""


################################################################################


def make_key_to_answer_dict(items):
    key_to_answer_dict = dict()
    for item in items:
        key = (item.firstname, item.lastname, item.h2, item.h3, item.question)
        assert key not in key_to_answer_dict, key
        key_to_answer_dict[key] = item.answer
    return key_to_answer_dict


def get_answer(*, firstname, lastname, h2, h3, question, key_to_answer_dict):
    key = (firstname, lastname, h2, h3, question)
    return key_to_answer_dict.get(key, " ")


################################################################################


def check_df(df):
    for (group, grouped_df) in df.groupby(by=["firstname", "lastname"]):
        for (block, questions) in survey_table.items():
            assert block in survey_table, block
            questions = set(survey_table[block])
            tags = set(grouped_df[grouped_df["h2"] == block]["question"])
            assert tags <= questions, (group, block, tags - questions)


################################################################################


def make_items(body):
    assert body is not None
    items = list()
    ##
    h1 = None
    h2 = None
    h3 = None
    firstname = None
    lastname = None
    for e in body.children:
        assert isinstance(e, Tag), e
        if e.name == "h1":
            assert is_pure(e)
            h1 = e.get_text()
            h2 = None
            h3 = None
            li = None
            # print(h1)
        elif e.name == "h2":
            assert is_pure(e)
            assert h1 is not None
            h2 = e.get_text()
            h3 = None
            li = None
            # print(h2)
        elif e.name == "h3":
            assert is_pure(e)
            assert h1 is not None
            assert h2 is not None
            h3 = e.get_text()
            li = None
            # print(h3)
        elif e.name == "ul":
            assert h1 is not None
            assert h2 is not None
            if h2 != "Données générales":
                assert h3 is not None, e
            ul = e
            for li in ul.children:
                assert isinstance(li, Tag)
                assert is_pure(li)
                assert li.name == "li"
                item = make_item_from_li(h1=h1, h2=h2, h3=h3, li=li)
                items.append(item)
        else:
            assert isinstance(e, Tag)
            assert e.name == "p"
            print(f"unexpected element: {e}")
    return items


def is_pure(e: PageElement):
    if isinstance(e, NavigableString):
        return True
    assert isinstance(e, Tag)
    for tag in e.find_all():
        if tag.name == "strong":
            continue
        raise AssertionError(f"Unexpected tag: {tag}")
    return True


################################################################################


def make_item_from_li(*, h1, h2, h3, li):
    ##
    parts = li.contents
    assert len(parts) == 2, parts
    # print(parts)
    ##
    question_part = parts[0]
    assert isinstance(question_part, Tag), question_part
    assert question_part.name == "strong", question_part
    question = question_part.get_text().strip()
    assert question.endswith(":"), question
    question = question[:-1]
    question = normalize(question)
    ##
    answer_part = parts[1]
    assert isinstance(answer_part, NavigableString), answer_part
    answer = answer_part.get_text()
    answer = normalize(answer)
    ##
    # print(question, "==>", answer)
    ##
    firstname, lastname = split_names(h1)
    item = make_item(
        firstname=firstname,
        lastname=lastname,
        h1=h1,
        h2=h2,
        h3=h3,
        question=question,
        answer=answer,
    )
    return item


def make_item(
    *,
    firstname,
    lastname,
    h1,
    h2,
    h3,
    question,
    answer,
):
    item = Item(
        firstname=firstname,
        lastname=lastname,
        h1=h1,
        h2=h2,
        h3=h3,
        question=question,
        answer=answer,
    )
    assert all(v is not None for (k, v) in item._asdict().items() if k != "h3"), item
    return item


################################################################################


def split_names(h1):
    assert h1 is not None
    name = h1.strip()
    assert name.count(",") == 1, name
    (lastname, firstname) = name.split(",")
    lastname = lastname.strip().title()
    firstname = firstname.strip().title()
    return (firstname, lastname)


def normalize(s):
    assert isinstance(s, str), s
    s = re.sub(r"\s+", " ", s).strip()
    return s


################################################################################

main()
