import os
import datetime
from unidecode import unidecode
from docx import Document
import docx
import lxml

# from docx.enum.style import WD_STYLE
from Library import survey
from Library import answers

## https://stackoverflow.com/a/63343007/7266382


def print_etree(n, x):
    print(n, "==>", lxml.etree.tostring(x, pretty_print=True).decode())


def set_language(*, lang, document):
    styles_element = document.styles.element
    rpr_default = styles_element.xpath("./w:docDefaults/w:rPrDefault/w:rPr")[0]
    # print_etree(1, rpr_default)
    lang_default = rpr_default.xpath("w:lang")[0]
    # print_etree(2, lang_default)
    lang_default.set(docx.oxml.shared.qn("w:val"), "fr-FR")
    # print_etree(3, lang_default)


def make_documents(*, key_to_answer_dict, persons):
    ##
    for (person_nb, (firstname, lastname)) in enumerate(persons, start=1):
        document = Document()
        set_language(lang="fr-FR", document=document)
        document.add_heading("Déclarations de liens d'intérêt", 0)
        now = datetime.datetime.now()
        now_string = s = now.strftime("%d/%m/%Y, %H:%M")
        document.add_paragraph(f"Date: {now_string}")
        assert document is not None
        print("Personne", (firstname, lastname))
        ##
        make_document_person(
            key_to_answer_dict=key_to_answer_dict,
            firstname=firstname,
            lastname=lastname,
            person_nb=person_nb,
            document=document,
        )
        ##
        filename = f"declaration_{lastname}_{firstname}.docx"
        filename = unidecode(filename).lower()
        dirname = "Declarations"
        os.makedirs(dirname, exist_ok=True)
        document.core_properties.language = "fr-FR"
        document.save(os.path.join(dirname, filename))


def make_document_person(
    *, key_to_answer_dict, firstname, lastname, person_nb, document
):
    document.add_heading(f"Personne {person_nb}: {lastname}, {firstname}", 1)
    for (block_nb, (block, item)) in enumerate(survey.structure.items(), start=1):
        # print("Bloc", block)
        document.add_heading(f"Bloc {block_nb}: {block}", 2)
        ##
        questions = item["questions"]
        description = item["description"].strip()
        nb_answers = item["nb_answers"]
        ##
        p = document.add_paragraph()
        p.add_run(description).italic = True
        ##
        for series_nb in range(1, nb_answers + 1):
            series = f"Réponse {series_nb}"
            document.add_heading(f"{series}", 3)
            prefix = (person_nb, block_nb, series_nb)
            make_document_questions(
                firstname=firstname,
                lastname=lastname,
                h2=block,
                h3=series,
                prefix=prefix,
                questions=questions,
                key_to_answer_dict=key_to_answer_dict,
                document=document,
            )
            prefix_tag = make_prefix_tag(prefix=prefix, suffix="9")
            add_validation_table(prefix_tag=prefix_tag, document=document)


def make_document_questions(
    *, firstname, lastname, h2, h3, prefix, questions, key_to_answer_dict, document
):
    for (question_nb, question) in enumerate(questions, start=1):
        answer = answers.get_answer(
            firstname=firstname,
            lastname=lastname,
            h2=h2,
            h3=h3,
            question=question,
            key_to_answer_dict=key_to_answer_dict,
        )
        p = document.add_paragraph(style="List Bullet")
        ## Number 9 is reserved for the validation question
        assert question_nb < 9
        prefix_tag = make_prefix_tag(prefix=prefix, suffix=question_nb)
        p.add_run(f"{prefix_tag} ")
        p.add_run(f"{question}:").bold = True
        p.add_run(f" {answer}")


def make_prefix_tag(*, prefix, suffix):
    full_prefix = prefix + (suffix,)
    prefix_string = ".".join((str(s) for s in full_prefix))
    return f"[{prefix_string}]"


def add_validation_table(*, prefix_tag, document):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[
        0
    ].text = f"{prefix_tag} Si le bloc est non vide, j'écris 'OK'\ndans la case de droite pour le valider"
    hdr_cells[1].text = ""
    table.allow_autofit = True
    table.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    table.columns[0].width = docx.shared.Cm(8)
    table.columns[1].width = docx.shared.Cm(1)
